"""Layout-preserving PDF translation — pdf2docx → translate paragraphs → LibreOffice → PDF.

Pro-only mode that preserves tables, columns, images, and general document structure.
"""

import logging
import os
import subprocess
import tempfile

from app.services.pdf_ai_service import (
    PdfAiError,
    _normalize_language_code,
    _language_label,
    _translate_document_text,
)

logger = logging.getLogger(__name__)

# LibreOffice binary path inside Docker
LIBREOFFICE_BIN = "libreoffice"
LIBREOFFICE_TIMEOUT = 120  # seconds


def _pdf_to_docx(input_pdf: str, output_docx: str) -> None:
    """Convert PDF to DOCX using pdf2docx, preserving layout."""
    try:
        from pdf2docx import Converter
    except ImportError as exc:
        raise PdfAiError(
            "Layout translation library is not installed on this server.",
            error_code="PDF2DOCX_NOT_INSTALLED",
            detail=str(exc),
        )

    try:
        cv = Converter(input_pdf)
        cv.convert(output_docx)
        cv.close()
    except Exception as exc:
        raise PdfAiError(
            "Failed to analyze PDF layout. The document structure may be too complex.",
            error_code="PDF2DOCX_CONVERSION_FAILED",
            detail=str(exc),
        )

    if not os.path.isfile(output_docx) or os.path.getsize(output_docx) < 100:
        raise PdfAiError(
            "Layout analysis produced an empty result.",
            error_code="PDF2DOCX_EMPTY_OUTPUT",
        )


def _translate_docx_paragraphs(
    docx_path: str,
    target_language: str,
    source_language: str | None = None,
    model_id: str | None = None,
) -> int:
    """Translate paragraph text inside a DOCX file in-place.

    Returns the number of paragraphs translated.
    """
    try:
        from docx import Document
    except ImportError as exc:
        raise PdfAiError(
            "DOCX processing library is not installed.",
            error_code="PYTHON_DOCX_NOT_INSTALLED",
            detail=str(exc),
        )

    doc = Document(docx_path)

    # Collect all non-empty paragraph texts
    paragraphs_to_translate: list[tuple[int, str]] = []
    for idx, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text and len(text) > 2:
            paragraphs_to_translate.append((idx, text))

    # Also collect table cell texts
    table_cells_to_translate: list[tuple[int, int, int, str]] = []
    for t_idx, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text and len(text) > 2:
                    table_cells_to_translate.append((t_idx, r_idx, c_idx, text))

    if not paragraphs_to_translate and not table_cells_to_translate:
        return 0

    # Batch all text for translation
    all_texts = [t for _, t in paragraphs_to_translate] + [
        t for _, _, _, t in table_cells_to_translate
    ]
    combined = "\n---PARA_SEP---\n".join(all_texts)

    translated = _translate_document_text(
        combined,
        target_language=target_language,
        source_language=source_language,
        model_id=model_id,
    )

    # Split back
    parts = translated["translation"].split("---PARA_SEP---")
    parts = [p.strip() for p in parts]

    # Apply translations to paragraphs
    para_count = len(paragraphs_to_translate)
    for i, (idx, _original) in enumerate(paragraphs_to_translate):
        if i < len(parts) and parts[i]:
            # Clear existing runs and set translated text
            para = doc.paragraphs[idx]
            if para.runs:
                # Preserve first run's formatting, clear others
                para.runs[0].text = parts[i]
                for run in para.runs[1:]:
                    run.text = ""
            else:
                para.text = parts[i]

    # Apply translations to table cells
    for j, (t_idx, r_idx, c_idx, _original) in enumerate(table_cells_to_translate):
        part_idx = para_count + j
        if part_idx < len(parts) and parts[part_idx]:
            cell = doc.tables[t_idx].rows[r_idx].cells[c_idx]
            if cell.paragraphs and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].text = parts[part_idx]
                for run in cell.paragraphs[0].runs[1:]:
                    run.text = ""
            else:
                cell.text = parts[part_idx]

    doc.save(docx_path)
    return len(paragraphs_to_translate) + len(table_cells_to_translate)


def _docx_to_pdf(docx_path: str, output_dir: str) -> str:
    """Convert DOCX to PDF using LibreOffice headless.

    Returns the path of the generated PDF.
    """
    try:
        result = subprocess.run(
            [
                LIBREOFFICE_BIN,
                "--headless",
                "--norestore",
                "--convert-to",
                "pdf",
                "--outdir",
                output_dir,
                docx_path,
            ],
            capture_output=True,
            text=True,
            timeout=LIBREOFFICE_TIMEOUT,
        )
    except FileNotFoundError:
        raise PdfAiError(
            "PDF conversion tool (LibreOffice) is not available on this server.",
            error_code="LIBREOFFICE_NOT_FOUND",
        )
    except subprocess.TimeoutExpired:
        raise PdfAiError(
            "PDF conversion timed out. The document may be too large.",
            error_code="LIBREOFFICE_TIMEOUT",
        )

    if result.returncode != 0:
        raise PdfAiError(
            "PDF conversion failed after translation.",
            error_code="LIBREOFFICE_CONVERSION_FAILED",
            detail=result.stderr[:500] if result.stderr else None,
        )

    # LibreOffice names the output using the input basename
    base = os.path.splitext(os.path.basename(docx_path))[0]
    pdf_path = os.path.join(output_dir, f"{base}.pdf")
    if not os.path.isfile(pdf_path):
        raise PdfAiError(
            "PDF conversion produced no output file.",
            error_code="LIBREOFFICE_NO_OUTPUT",
        )
    return pdf_path


def translate_pdf_layout(
    input_path: str,
    target_language: str,
    output_path: str,
    original_filename: str,
    source_language: str | None = None,
    model_id: str | None = None,
) -> dict:
    """Full layout-preserving translation pipeline.

    Returns:
        {
            "paragraphs_translated": int,
            "pages": int,
            "target_language": str,
            "provider": str,
        }
    """
    normalized_target = _normalize_language_code(target_language)
    normalized_source = _normalize_language_code(source_language, default="auto")

    with tempfile.TemporaryDirectory(prefix="translate_layout_") as tmpdir:
        docx_path = os.path.join(tmpdir, "source.docx")

        # Step 1: PDF → DOCX (layout preserved)
        _pdf_to_docx(input_path, docx_path)

        # Step 2: Translate paragraphs in-place
        count = _translate_docx_paragraphs(
            docx_path,
            target_language=normalized_target,
            source_language=normalized_source,
            model_id=model_id,
        )

        if count == 0:
            raise PdfAiError(
                "No translatable text found in the document layout.",
                error_code="LAYOUT_NO_TEXT",
            )

        # Step 3: DOCX → PDF via LibreOffice
        generated_pdf = _docx_to_pdf(docx_path, tmpdir)

        # Copy to final output path
        import shutil
        shutil.copy2(generated_pdf, output_path)

    from PyPDF2 import PdfReader

    pages = len(PdfReader(input_path).pages)

    return {
        "paragraphs_translated": count,
        "pages": pages,
        "target_language": normalized_target,
        "source_language": normalized_source,
        "provider": "layout",
    }
